from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "competition" / "proposal.md"
DOCX_OUT = ROOT / "docs" / "competition" / "MoonSeal项目申报书.docx"
PDF_OUT = ROOT / "docs" / "competition" / "MoonSeal项目申报书.pdf"


def parse_markdown():
    sections = []
    current = None
    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            current = {"level": 1, "title": line[2:], "body": []}
            sections.append(current)
        elif line.startswith("## "):
            current = {"level": 2, "title": line[3:], "body": []}
            sections.append(current)
        elif current is not None:
            current["body"].append(line)
    return sections


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Microsoft YaHei"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(31, 77, 121)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)


def add_metadata_table(doc):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目名称", "MoonSeal"),
        ("项目方向", "软件分析框架 / 工程质量工具"),
        ("GitLink 仓库", "https://gitlink.org.cn/LL1266/moonseal"),
        ("GitHub 仓库", "https://github.com/LL728/L28L"),
        ("开源协议", "Apache-2.0"),
    ]
    for row, (key, value) in zip(table.rows, rows):
        left, right = row.cells
        left.text = ""
        right.text = ""
        lrun = left.paragraphs[0].add_run(key)
        rrun = right.paragraphs[0].add_run(value)
        set_run_font(lrun, bold=True)
        set_run_font(rrun)
        shade_cell(left, "EAF1F7")
    doc.add_paragraph()


def add_body_line(doc, line):
    if line.startswith("- "):
        doc.add_paragraph(line[2:], style="List Bullet")
    elif len(line) > 3 and line[:3] in {f"{i}. " for i in range(1, 10)}:
        doc.add_paragraph(line[3:], style="List Number")
    elif line.strip():
        p = doc.add_paragraph()
        p.add_run(line)


def build_docx(sections):
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("MoonSeal 项目申报书"), size=20, bold=True, color=(31, 77, 121))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("MoonBit 测试充分性与发布质量门禁工具"), size=11)

    add_metadata_table(doc)

    for item in sections:
        if item["level"] == 1:
            continue
        if item["title"] == "仓库地址":
            continue
        doc.add_heading(item["title"], level=1)
        in_code = False
        code_lines = []
        for line in item["body"]:
            if line.startswith("```"):
                if in_code:
                    p = doc.add_paragraph("\n".join(code_lines))
                    for run in p.runs:
                        set_run_font(run, name="Consolas", size=9)
                    code_lines = []
                in_code = not in_code
            elif in_code:
                code_lines.append(line)
            else:
                add_body_line(doc, line)

    doc.save(DOCX_OUT)


def register_font():
    for font_path in [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    ]:
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("CNFont", str(font_path)))
            return "CNFont"
    return "Helvetica"


def esc(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(sections):
    font = register_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontName=font,
        fontSize=20,
        leading=26,
        alignment=1,
        textColor=colors.HexColor("#1f4d79"),
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleStyle",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=10.5,
        leading=15,
        alignment=1,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyStyle",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=10,
        leading=14,
        spaceAfter=5,
    )
    heading_style = ParagraphStyle(
        "HeadingStyle",
        parent=styles["Heading1"],
        fontName=font,
        fontSize=12.5,
        leading=17,
        textColor=colors.HexColor("#1f4d79"),
        spaceBefore=8,
        spaceAfter=5,
    )
    code_style = ParagraphStyle(
        "CodeStyle",
        parent=body_style,
        fontName="Courier",
        fontSize=8,
        leading=10,
        leftIndent=8,
    )

    story = [
        Paragraph("MoonSeal 项目申报书", title_style),
        Paragraph("MoonBit 测试充分性与发布质量门禁工具", subtitle_style),
        Spacer(1, 4),
    ]
    table = Table(
        [
            ["项目名称", "MoonSeal"],
            ["项目方向", "软件分析框架 / 工程质量工具"],
            ["GitLink 仓库", "https://gitlink.org.cn/LL1266/moonseal"],
            ["GitHub 仓库", "https://github.com/LL728/L28L"],
            ["开源协议", "Apache-2.0"],
        ],
        colWidths=[32 * mm, 128 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eaf1f7")),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#9aaec3")),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.extend([table, Spacer(1, 7)])

    for item in sections:
        if item["level"] == 1 or item["title"] == "仓库地址":
            continue
        story.append(Paragraph(esc(item["title"]), heading_style))
        in_code = False
        code_lines = []
        for line in item["body"]:
            if line.startswith("```"):
                if in_code:
                    story.append(Paragraph("<br/>".join(esc(x) for x in code_lines), code_style))
                    code_lines = []
                in_code = not in_code
            elif in_code:
                code_lines.append(line)
            elif line.startswith("- "):
                story.append(Paragraph("• " + esc(line[2:]), body_style))
            elif len(line) > 3 and line[:3] in {f"{i}. " for i in range(1, 10)}:
                story.append(Paragraph(esc(line), body_style))
            elif line.strip():
                story.append(Paragraph(esc(line), body_style))

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=22 * mm,
        leftMargin=22 * mm,
        topMargin=22 * mm,
        bottomMargin=22 * mm,
    )
    doc.build(story)


def main():
    sections = parse_markdown()
    build_docx(sections)
    build_pdf(sections)
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()
